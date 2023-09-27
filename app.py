import os
import uuid
import shutil
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from datetime import datetime, timedelta

# Função para calcular a data de saída com base na data inicial e no número de meses
def calcular_data_saida(data_inicial, meses):
    data_formato = "%d/%m/%Y"
    data_inicial = datetime.strptime(data_inicial, data_formato)
    data_saida = data_inicial + timedelta(days=30 * int(meses))
    return data_saida.strftime(data_formato)

# Função para gerar o contrato
def gerar_contrato(valor, meses, data_inicial):
    global resultado_simulacao_label

    try:
        valor = float(valor)
        meses = int(meses)
        data_saida = calcular_data_saida(data_inicial, meses)

        # Determine qual modelo de contrato carregar com base na seleção do usuário
        if tipo_pessoa_selecionado == "Física":
            modelo_contrato = Document("contrato_pessoa_fisica.docx")
        elif tipo_pessoa_selecionado == "Jurídica":
            modelo_contrato = Document("contrato_pessoa_juridica.docx")
        else:
            # Trate qualquer outro caso, se necessário
            modelo_contrato = None

        if modelo_contrato:
            # Dicionário de tags e valores a serem substituídos
            tags_e_valores = {
                "<NOME_LOCATARIO>": nome_entry.get(),
                "<CPF>": cpf_entry.get(),
                "<RG>": rg_entry.get(),
                "<VALOR_ALUGUEL>": f'R${valor:.2f} por mês',
                "<DATA_INICIAL>": data_inicial,
                "<DATA_SAIDA>": data_saida,
                "<LOGRADOURO>": logradouro_entry.get(),
                "<NUMERO>": numero_entry.get(),
                "<COMPLEMENTO>": complemento_entry.get(),
                "<BAIRRO>": bairro_entry.get(),
                "<ESTADO>": estado_entry.get(),
                "<CIDADE>": cidade_entry.get(),
                "<CEP>": cep_entry.get(),
                "<EMAIL>": email_entry.get(),
                # Adicione mais tags e valores conforme necessário
            }

            # Substitua as tags pelo valor correspondente no modelo
            for paragrafo in modelo_contrato.paragraphs:
                for tag, valor in tags_e_valores.items():
                    if tag in paragrafo.text:
                        paragrafo.text = paragrafo.text.replace(tag, valor)

            # Gerar um nome de arquivo único e aleatório usando UUID
            nome_arquivo = str(uuid.uuid4())[:8] + '_contrato_aluguel.docx'

            # Salvar o contrato gerado em um arquivo
            modelo_contrato.save(nome_arquivo)

            # Obter o diretório de downloads do usuário
            pasta_downloads = os.path.expanduser("~\\Downloads")

            # Mover o arquivo para a pasta de downloads
            shutil.move(nome_arquivo, os.path.join(pasta_downloads, nome_arquivo))

            resultado_simulacao_label.config(text="Contrato gerado com sucesso. Verifique a pasta de downloads.")

        else:
            resultado_simulacao_label.config(text="Selecione o tipo de pessoa (Física ou Jurídica) antes de gerar o contrato.")

    except ValueError:
        resultado_simulacao_label.config(text="Por favor, insira valores válidos.")



# Função para realizar o login
def fazer_login():
    usuario = entry_usuario.get()
    senha = entry_senha.get()

    if usuario in dados_login and senha == dados_login[usuario]:
        # Login bem-sucedido
        login_frame.pack_forget()
        home_frame.pack()

    else:
        messagebox.showerror("Erro de Login", "Login falhou. Verifique usuário e senha.")

# Função para exibir campos de pessoa física
def mostrar_campos_pessoa_fisica():
    global tipo_pessoa_selecionado
    tipo_pessoa_selecionado = "Física"
    campos_pessoa_juridica_frame.pack_forget()
    campos_pessoa_fisica_frame.pack()

# Função para exibir campos de pessoa jurídica
def mostrar_campos_pessoa_juridica():
    global tipo_pessoa_selecionado
    tipo_pessoa_selecionado = "Jurídica"
    campos_pessoa_fisica_frame.pack_forget()
    campos_pessoa_juridica_frame.pack()

# Função para realizar a simulação de aluguel
def simular_aluguel(valor, meses, data_inicial):
    global resultado_simulacao_label

    if not valor or not meses or not data_inicial:
        resultado_simulacao_label.config(text="Por favor, preencha todos os campos.", fg="red")
        # Oculta o botão "Gerar Contrato" quando ocorre um erro
        gerar_contrato_button.pack_forget()
        return

    try:
        valor = float(valor)
        meses = int(meses)
        data_formato = "%d/%m/%Y"
        data_inicio = datetime.strptime(data_inicial, data_formato)
        data_saida = calcular_data_saida(data_inicial, meses)
        valor_total = valor * meses
        resultado = f"Simulação de aluguel realizada com sucesso!\nValor do aluguel: R${valor:.2f}\nData de Entrada: {data_inicial}\nData de Saída: {data_saida}\nTotal Pago: R${valor_total:.2f}"
        resultado_simulacao_label.config(text=resultado, fg="black")
        
        # Exibe o botão "Gerar Contrato" quando a simulação é bem-sucedida
        gerar_contrato_button.pack(side="top", padx=10, pady=10)
    except ValueError:
        resultado_simulacao_label.config(text="Por favor, insira valores válidos.", fg="red")

        # Oculta o botão "Gerar Contrato" quando ocorre um erro
        gerar_contrato_button.pack_forget()

# Dados de login de exemplo
dados_login = {"admin": "admin"}

# Configuração da janela principal
root = tk.Tk()
root.title("Sistema de Simulação de Aluguel!")

# Maximize a janela
root.state('zoomed')

# Frame de login
login_frame = tk.Frame(root)
login_frame.pack()

# Crie um estilo
style = ttk.Style()
style.configure("BoldLabel.TLabel", font=("Helvetica", 16, "bold"))
style.configure("Green.TButton", background="green", foreground="black", font=("Helvetica", 12, "bold"))

# Frame de login
login_frame = tk.Frame()

label_usuario = ttk.Label(login_frame, text="Usuário", style="BoldLabel.TLabel")
label_usuario.pack(pady=(20, 5))

entry_usuario = tk.Entry(login_frame, width=40, borderwidth=2, font=("Helvetica", 14))
entry_usuario.pack()

label_senha = ttk.Label(login_frame, text="Senha", style="BoldLabel.TLabel")
label_senha.pack(pady=(20, 5))

entry_senha = tk.Entry(login_frame, width=40, borderwidth=2, font=("Helvetica", 14), show="*")  
entry_senha.pack()

botao_login = ttk.Button(login_frame, text="Login", command=fazer_login, style="Green.TButton")
botao_login.pack(pady=20)

login_frame.pack()

# Frame da tela principal (home)
home_frame = tk.Frame(root)

label_selecione = tk.Label(home_frame, text="Selecione uma opção:", font=("Helvetica", 16, "bold"))
label_selecione.pack(pady=10)

botao_pessoa_fisica = tk.Button(home_frame, text="Pessoa Física", command=mostrar_campos_pessoa_fisica)
botao_pessoa_fisica.pack(side="top", padx=10, pady=10)

botao_pessoa_juridica = tk.Button(home_frame, text="Pessoa Jurídica", command=mostrar_campos_pessoa_juridica)
botao_pessoa_juridica.pack(side="top", padx=10, pady=10)

# Campos para entrada de dados do aluguel
valor_aluguel_label = tk.Label(home_frame, text="Valor do Aluguel:", font=("Helvetica", 12, "bold"))
valor_aluguel_label.pack()
valor_aluguel_entry = tk.Entry(home_frame, width=30, borderwidth=2, font=("Helvetica", 14))
valor_aluguel_entry.pack()

meses_label = tk.Label(home_frame, text="Número de Meses:", font=("Helvetica", 12, "bold"))
meses_label.pack()
meses_entry = tk.Entry(home_frame, width=30, borderwidth=2, font=("Helvetica", 14))
meses_entry.pack()

data_inicial_label = tk.Label(home_frame, text="Data Inicial:", font=("Helvetica", 12, "bold"))
data_inicial_label.pack()
data_inicial_entry = tk.Entry(home_frame, width=30, borderwidth=2, font=("Helvetica", 14))
data_inicial_entry.pack()

# Frame para campos de pessoa física
campos_pessoa_fisica_frame = tk.Frame(home_frame)

# Campos de entrada de dados de pessoa física
nome_label = tk.Label(campos_pessoa_fisica_frame, text="Nome:", font=("Helvetica", 10, "bold"))
nome_label.grid(row=0, column=0, padx=10, pady=5)
nome_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
nome_entry.grid(row=0, column=1, padx=10, pady=5)

cpf_label = tk.Label(campos_pessoa_fisica_frame, text="CPF:", font=("Helvetica", 10, "bold"))
cpf_label.grid(row=0, column=2, padx=10, pady=5)
cpf_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
cpf_entry.grid(row=0, column=3, padx=10, pady=5)

rg_label = tk.Label(campos_pessoa_fisica_frame, text="RG:", font=("Helvetica", 10, "bold"))
rg_label.grid(row=1, column=0, padx=10, pady=5)
rg_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
rg_entry.grid(row=1, column=1, padx=10, pady=5)

orgao_emissor_label = tk.Label(campos_pessoa_fisica_frame, text="Orgão Emissor/UF:", font=("Helvetica", 10, "bold"))
orgao_emissor_label.grid(row=1, column=0, padx=10, pady=5)
orgao_emissor_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
orgao_emissor_entry.grid(row=1, column=1, padx=10, pady=5)

data_nascimento_label = tk.Label(campos_pessoa_fisica_frame, text="Data de Nascimento:", font=("Helvetica", 10, "bold"))
data_nascimento_label.grid(row=2, column=0, padx=10, pady=5)
data_nascimento_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
data_nascimento_entry.grid(row=2, column=1, padx=10, pady=5)

estado_civil_label = tk.Label(campos_pessoa_fisica_frame, text="Estado Civil:", font=("Helvetica", 10, "bold"))
estado_civil_label.grid(row=2, column=2, padx=10, pady=5)
estado_civil_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
estado_civil_entry.grid(row=2, column=3, padx=10, pady=5)

logradouro_label = tk.Label(campos_pessoa_fisica_frame, text="Logradouro:", font=("Helvetica", 10, "bold"))
logradouro_label.grid(row=3, column=0, padx=10)
logradouro_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
logradouro_entry.grid(row=3, column=1)

numero_label = tk.Label(campos_pessoa_fisica_frame, text="Número:", font=("Helvetica", 10, "bold"))
numero_label.grid(row=3, column=2, padx=10)
numero_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
numero_entry.grid(row=3, column=3)

complemento_label = tk.Label(campos_pessoa_fisica_frame, text="Complemento:", font=("Helvetica", 10, "bold"))
complemento_label.grid(row=3, column=4, padx=10)
complemento_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
complemento_entry.grid(row=3, column=5)

cep_label = tk.Label(campos_pessoa_fisica_frame, text="CEP:", font=("Helvetica", 10, "bold"))
cep_label.grid(row=4, column=0, padx=10)
cep_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
cep_entry.grid(row=4, column=1)

bairro_label = tk.Label(campos_pessoa_fisica_frame, text="Bairro:", font=("Helvetica", 10, "bold"))
bairro_label.grid(row=4, column=2, padx=10)
bairro_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
bairro_entry.grid(row=4, column=3)

# Lista Estado
listEstados = ["RJ", "SP", "ES", "MG", "PR", "SC", "RS", "MS", "GO", "AC", "AL", "AP", "AM", "BA", "CE", "DF", "MA", "MT", "PA", "PB", "PE", "PI", "RN", "RO", "RR", "SE", "TO"]

estado_label = tk.Label(campos_pessoa_fisica_frame, text="Estado:", font=("Helvetica", 10, "bold"))
estado_label.grid(row=5, column=0, padx=10)
estado_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
estado_entry.grid(row=5, column=1)

# Lista Cidade
listCidades = []

cidade_label = tk.Label(campos_pessoa_fisica_frame, text="Cidade:", font=("Helvetica", 10, "bold"))
cidade_label.grid(row=5, column=2, padx=10)
cidade_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
cidade_entry.grid(row=5, column=3)

telefone_label = tk.Label(campos_pessoa_fisica_frame, text="Telefone:", font=("Helvetica", 10, "bold"))
telefone_label.grid(row=6, column=0, padx=10, pady=5)
telefone_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
telefone_entry.grid(row=6, column=1, padx=10, pady=5)

email_label = tk.Label(campos_pessoa_fisica_frame, text="Email:", font=("Helvetica", 10, "bold"))
email_label.grid(row=7, column=0, padx=10, pady=5)
email_entry = tk.Entry(campos_pessoa_fisica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
email_entry.grid(row=7, column=1, padx=10, pady=5)

# Frame para campos de pessoa jurídica
campos_pessoa_juridica_frame = tk.Frame(home_frame)

# Campos de entrada de dados de pessoa jurídica
razao_social_label = tk.Label(campos_pessoa_juridica_frame, text="Razão Social:", font=("Helvetica", 10, "bold"))
razao_social_label.grid(row=0, column=0, padx=10, pady=5)
razao_social_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
razao_social_entry.grid(row=0, column=1, padx=10, pady=5)

nome_fantasia_label = tk.Label(campos_pessoa_juridica_frame, text="Nome Fantasia:", font=("Helvetica", 10, "bold"))
nome_fantasia_label.grid(row=0, column=2, padx=10, pady=5)
nome_fantasia_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
nome_fantasia_entry.grid(row=0, column=3, padx=10, pady=5)

cnpj_label = tk.Label(campos_pessoa_juridica_frame, text="CNPJ:", font=("Helvetica", 10, "bold"))
cnpj_label.grid(row=1, column=0, padx=10, pady=5)
cnpj_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
cnpj_entry.grid(row=1, column=1, padx=10, pady=5)

data_abertura_label = tk.Label(campos_pessoa_juridica_frame, text="Data de Abertura:", font=("Helvetica", 10, "bold"))
data_abertura_label.grid(row=1, column=2, padx=10, pady=5)
data_abertura_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
data_abertura_entry.grid(row=1, column=3, padx=10, pady=5)

logradouro_label = tk.Label(campos_pessoa_juridica_frame, text="Logradouro:", font=("Helvetica", 10, "bold"))
logradouro_label.grid(row=2, column=0, padx=10)
logradouro_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
logradouro_entry.grid(row=2, column=1)

numero_label = tk.Label(campos_pessoa_juridica_frame, text="Número:", font=("Helvetica", 10, "bold"))
numero_label.grid(row=2, column=2, padx=10)
numero_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
numero_entry.grid(row=2, column=3)

complemento_label = tk.Label(campos_pessoa_juridica_frame, text="Complemento:", font=("Helvetica", 10, "bold"))
complemento_label.grid(row=2, column=4, padx=10)
complemento_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
complemento_entry.grid(row=2, column=5)

cep_label = tk.Label(campos_pessoa_juridica_frame, text="CEP:", font=("Helvetica", 10, "bold"))
cep_label.grid(row=3, column=0, padx=10)
cep_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
cep_entry.grid(row=3, column=1)

bairro_label = tk.Label(campos_pessoa_juridica_frame, text="Bairro:", font=("Helvetica", 10, "bold"))
bairro_label.grid(row=3, column=2, padx=10)
bairro_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
bairro_entry.grid(row=3, column=3)

# Lista Estado
listEstados = ["RJ", "SP", "ES", "MG", "PR", "SC", "RS", "MS", "GO", "AC", "AL", "AP", "AM", "BA", "CE", "DF", "MA", "MT", "PA", "PB", "PE", "PI", "RN", "RO", "RR", "SE", "TO"]

estado_label = tk.Label(campos_pessoa_juridica_frame, text="Estado:", font=("Helvetica", 10, "bold"))
estado_label.grid(row=4, column=0, padx=10)
estado_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
estado_entry.grid(row=4, column=1)

# Lista Cidade
listCidades = []

cidade_label = tk.Label(campos_pessoa_juridica_frame, text="Cidade:", font=("Helvetica", 10, "bold"))
cidade_label.grid(row=4, column=2, padx=10)
cidade_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
cidade_entry.grid(row=4, column=3)

telefone_empresa_label = tk.Label(campos_pessoa_juridica_frame, text="Telefone da Empresa:", font=("Helvetica", 10, "bold"))
telefone_empresa_label.grid(row=5, column=0, padx=10, pady=5)
telefone_empresa_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
telefone_empresa_entry.grid(row=5, column=1, padx=10, pady=5)

email_empresa_label = tk.Label(campos_pessoa_juridica_frame, text="Email da Empresa:", font=("Helvetica", 10, "bold"))
email_empresa_label.grid(row=6, column=0, padx=10, pady=5)
email_empresa_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
email_empresa_entry.grid(row=6, column=1, padx=10, pady=5)

nome_socio_label = tk.Label(campos_pessoa_juridica_frame, text="Nome do Sócio:", font=("Helvetica", 10, "bold"))
nome_socio_label.grid(row=7, column=0, padx=10, pady=5)
nome_socio_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
nome_socio_entry.grid(row=7, column=1, padx=10, pady=5)

cpf_socio_label = tk.Label(campos_pessoa_juridica_frame, text="CPF do Sócio:", font=("Helvetica", 10, "bold"))
cpf_socio_label.grid(row=7, column=2, padx=10, pady=5)
cpf_socio_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
cpf_socio_entry.grid(row=7, column=3, padx=10, pady=5)

rg_socio_label = tk.Label(campos_pessoa_juridica_frame, text="RG do Sócio:", font=("Helvetica", 10, "bold"))
rg_socio_label.grid(row=8, column=0, padx=10, pady=5)
rg_socio_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
rg_socio_entry.grid(row=8, column=1, padx=10, pady=5)

orgao_emissor_label = tk.Label(campos_pessoa_juridica_frame, text="Orgão Emissor/UF:", font=("Helvetica", 10, "bold"))
orgao_emissor_label.grid(row=8, column=2, padx=10, pady=5)
orgao_emissor_entry = tk.Entry(campos_pessoa_juridica_frame, width=30, borderwidth=2, font=("Helvetica", 14))
orgao_emissor_entry.grid(row=8, column=3, padx=10, pady=5)

# Botões para simular e gerar contrato
simular_button = tk.Button(home_frame, text="Simular", command=lambda: simular_aluguel(valor_aluguel_entry.get(), meses_entry.get(), data_inicial_entry.get()))
simular_button.pack(side="top", padx=10, pady=10)

gerar_contrato_button = tk.Button(home_frame, text="Gerar Contrato", command=lambda: gerar_contrato(valor_aluguel_entry.get(), meses_entry.get(), data_inicial_entry.get()))
gerar_contrato_button.pack(side="top", padx=10, pady=10)
gerar_contrato_button.pack_forget()  # Oculta o botão inicialmente

# Rótulo para exibir o resultado da simulação
resultado_simulacao_label = tk.Label(home_frame, text="Aluguel não Simulada!", font=("Helvetica", 12, "bold"))
resultado_simulacao_label.pack()

# Oculta os campos inicialmente
campos_pessoa_fisica_frame.pack_forget()
campos_pessoa_juridica_frame.pack_forget()

# Oculta a tela principal inicialmente
home_frame.pack_forget()

root.mainloop()
